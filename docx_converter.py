"""
NPvert DOCX Converter
Converts DOCX files to LaTeX format.
"""

from docx import Document
from typing import Optional
import re
import os


class DOCXConverter:
    """Converts DOCX documents to LaTeX."""
    
    def __init__(self):
        self.latex_output = []
        
    def convert(self, file_path: str) -> str:
        """Convert DOCX file to LaTeX string."""
        doc = Document(file_path)
        self.latex_output = []
        
        # Add document class
        self.latex_output.append("\\documentclass[12pt,a4paper]{article}")
        self.latex_output.append("\\usepackage[utf8]{inputenc}")
        self.latex_output.append("\\usepackage{graphicx}")
        self.latex_output.append("\\usepackage{booktabs}")
        self.latex_output.append("\\usepackage{amsmath}")
        self.latex_output.append("\\usepackage{hyperref}")
        self.latex_output.append("")
        self.latex_output.append("\\begin{document}")
        self.latex_output.append("")
        
        # Process paragraphs
        for para in doc.paragraphs:
            latex_line = self._convert_paragraph(para)
            if latex_line:
                self.latex_output.append(latex_line)
        
        # Process tables
        for table in doc.tables:
            latex_table = self._convert_table(table)
            if latex_table:
                self.latex_output.append(latex_table)
        
        self.latex_output.append("")
        self.latex_output.append("\\end{document}")
        
        return '\n'.join(self.latex_output)
    
    def _convert_paragraph(self, para) -> Optional[str]:
        """Convert a paragraph to LaTeX."""
        text = para.text.strip()
        if not text:
            return None
        
        # Escape special characters
        text = self._escape_latex(text)
        
        # Check style
        style_name = para.style.name.lower() if para.style else ""
        
        if 'heading 1' in style_name or 'title' in style_name:
            return f"\\section{{{text}}}"
        elif 'heading 2' in style_name:
            return f"\\subsection{{{text}}}"
        elif 'heading 3' in style_name:
            return f"\\subsubsection{{{text}}}"
        elif 'abstract' in style_name:
            return f"\\begin{{abstract}}\n{text}\n\\end{{abstract}}"
        else:
            # Check for bold/italic in runs
            formatted = ""
            for run in para.runs:
                run_text = self._escape_latex(run.text)
                if run.bold and run.italic:
                    formatted += f"\\textbf{{\\textit{{{run_text}}}}}"
                elif run.bold:
                    formatted += f"\\textbf{{{run_text}}}"
                elif run.italic:
                    formatted += f"\\textit{{{run_text}}}"
                else:
                    formatted += run_text
            return formatted if formatted else text
    
    def _convert_table(self, table) -> str:
        """Convert a table to LaTeX."""
        if not table.rows:
            return ""
        
        num_cols = len(table.rows[0].cells)
        col_spec = "c" * num_cols
        
        lines = ["\\begin{table}[h]", "\\centering", f"\\begin{{tabular}}{{{col_spec}}}"]
        lines.append("\\toprule")
        
        for i, row in enumerate(table.rows):
            cells = [self._escape_latex(cell.text.strip()) for cell in row.cells]
            lines.append(" & ".join(cells) + " \\\\")
            if i == 0:
                lines.append("\\midrule")
        
        lines.append("\\bottomrule")
        lines.append("\\end{tabular}")
        lines.append("\\caption{Table converted from DOCX}")
        lines.append("\\end{table}")
        
        return '\n'.join(lines)
    
    def _escape_latex(self, text: str) -> str:
        """Escape LaTeX special characters."""
        text = text.replace('\\', '\\textbackslash ')
        text = text.replace('&', '\\&')
        text = text.replace('%', '\\%')
        text = text.replace('$', '\\$')
        text = text.replace('#', '\\#')
        text = text.replace('_', '\\_')
        text = text.replace('{', '\\{')
        text = text.replace('}', '\\}')
        text = text.replace('~', '\\textasciitilde ')
        text = text.replace('^', '\\textasciicircum ')
        return text
    
    def convert_to_file(self, docx_path: str, output_path: str):
        """Convert DOCX and save to file."""
        latex = self.convert(docx_path)
        with open(output_path, 'w', encoding='utf-8') as f:
            f.write(latex)
        return output_path


def is_docx(file_path: str) -> bool:
    """Check if file is a DOCX."""
    return file_path.lower().endswith('.docx')


def read_input_file(file_path: str) -> str:
    """Read input file, converting if necessary."""
    if is_docx(file_path):
        converter = DOCXConverter()
        return converter.convert(file_path)
    else:
        with open(file_path, 'r', encoding='utf-8', errors='ignore') as f:
            return f.read()
